import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create(name_of_file: str, words: dict):
    nums = 1
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(15)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Dictionary")
    run.bold = True

    for k, v in words.items():
        text = f"{nums}) Word: {k}"
        for key, value in v[0].items():
            text += '\n' + key + ': '
            if type(value) is list:
                for j in value:
                    text += '\n\t' + j
                text += '\n'
            else:
                text += value + '\n'
        txt = document.add_paragraph()
        txt.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        txt.add_run(text)
        nums += 1

    document.save(f'{name_of_file}.docx')


def delete1(word):
    location = "C:/Users/abdul/PycharmProjects/PyDict_TeleBot/Telebot/"
    path = os.path.join(location, word)
    os.remove(path)
